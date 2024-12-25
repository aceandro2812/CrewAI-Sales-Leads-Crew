from docx import Document
from crewai import Agent, Task, Crew , Process ,LLM
import os
from typing import List, Dict
from pydantic import BaseModel, Field

gemini_llm  = LLM(
    model="openai/gemini-2.0-flash-exp",
    temperature=0.7,
    base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
    api_key="AIzaSyA6Gd_kJL0g8XCMZXJ-uJwbTDYcac1zqGk"
)

def write_to_docx(filename, documentation):
    doc = Document()
    doc.add_heading(f'Documentation for {filename}', level=1)
    doc.add_paragraph(documentation)
    doc.save(f'{filename}_documentation.docx')

from crewai_tools import BaseTool
from typing import Optional

class CodeReadTool(BaseTool):
    name: str = "Read Code File"
    description: str = "Reads and returns the content of a specified code file."

    def _run(
        self,
        file_path: str,
        **kwargs
    ) -> str:
        """Reads a file and returns its content."""
        try:
            with open(file_path, 'r') as f:
                content = f.read()
            return content
        except FileNotFoundError:
            return f"Error: File '{file_path}' not found."
        except Exception as e:
            return f"Error reading file '{file_path}': {e}"

code_read_tool = CodeReadTool()

directory_analyzer = Agent(
    role='Directory Analyzer',
    goal='Analyze the directory structure and identify all code files.',
    backstory="""You are an expert in navigating and understanding complex code repositories.
    You can identify all relevant code files in a directory.""",
    llm=gemini_llm,
    verbose=True,
    allow_delegation=True
)

code_reader = Agent(
    role='Code Reader',
    goal='Read and understand the content of individual code files.',
    backstory="""You are a skilled code analyst capable of understanding the logic,
    purpose, and functionality of any given code file.""",
    llm=gemini_llm,
    tools=[code_read_tool],
    verbose=True,
    allow_delegation=True
)

documentation_writer = Agent(
    role='Documentation Writer',
    goal='Generate comprehensive documentation for each code file and write it to a Word document.',
    backstory="""You are a technical writer with a talent for explaining complex
    technical details in a clear, concise, and understandable manner.""",
    llm=gemini_llm,
    verbose=True,
    allow_delegation=False
)

analyze_directory_task = Task(
    description='Analyze the project directory to identify all code files.',
    agent=directory_analyzer
)

read_code_task = Task(
    description='Read the content of each identified code file.',
    agent=code_reader
)

write_documentation_task = Task(
    description='Generate documentation for each code file and save it as a Word document.',
    agent=documentation_writer
)

crew = Crew(
    agents=[directory_analyzer, code_reader, documentation_writer],
    tasks=[analyze_directory_task, read_code_task, write_documentation_task],
    process=Process.sequential,
    verbose=2
)

result = crew.kickoff()
